import io
import os
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel

from checker_core import verify_grant_text

load_dotenv()

app = FastAPI(title="GrantChecker API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:8000", "http://127.0.0.1:8000", "http://127.0.0.1:3000"],
    allow_methods=["*"],
    allow_headers=["*"],
)

AVAILABLE_MODELS = {
    "anthropic": ["claude-sonnet-4-6", "claude-opus-4-6", "claude-haiku-4-5-20251001"],
    "yandex": ["yandexgpt/latest", "yandexgpt-lite/latest"],
}

SECTION_LABELS = {
    "critical":    "Критические ошибки",
    "significant": "Существенные замечания",
    "minor":       "Незначительные замечания",
    "confirmed":   "Подтверждённые факты",
    "needs_manual": "Требует ручной проверки",
}

# Heading colours for DOCX sections
SECTION_COLORS = {
    "critical":    RGBColor(0x72, 0x1C, 0x24),
    "significant": RGBColor(0x85, 0x64, 0x04),
    "minor":       RGBColor(0x0C, 0x54, 0x60),
    "confirmed":   RGBColor(0x15, 0x57, 0x24),
    "needs_manual": RGBColor(0x38, 0x3D, 0x41),
}


class VerifyRequest(BaseModel):
    text: str
    provider: str = "anthropic"
    model: str = "claude-sonnet-4-6"


class ExportItem(BaseModel):
    quote: str = ""
    issue: str = ""
    recommendation: str = ""


class ExportReport(BaseModel):
    critical: list[ExportItem] = []
    significant: list[ExportItem] = []
    minor: list[ExportItem] = []
    confirmed: list[ExportItem] = []
    needs_manual: list[ExportItem] = []


class ExportRequest(BaseModel):
    text: str
    report: ExportReport
    provider: str = "anthropic"
    model: str = "claude-sonnet-4-6"


def _build_docx(req: ExportRequest) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading("GrantChecker — Отчёт о проверке", level=1)
    title.runs[0].font.size = Pt(18)

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").italic = True
    meta.add_run(f"   Провайдер: {req.provider} / {req.model}").italic = True

    doc.add_paragraph()

    # Original text
    doc.add_heading("Исходный текст", level=2)
    original = doc.add_paragraph(req.text)
    original.style.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("Результаты проверки", level=2)

    # Sections
    for key, label in SECTION_LABELS.items():
        items = getattr(req.report, key, [])

        heading = doc.add_heading(label, level=3)
        color = SECTION_COLORS.get(key, RGBColor(0, 0, 0))
        for run in heading.runs:
            run.font.color.rgb = color

        if not items:
            doc.add_paragraph("Замечаний нет").italic = True
        else:
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                if item.quote:
                    run = p.add_run(f"«{item.quote}» ")
                    run.italic = True
                if item.issue:
                    p.add_run(item.issue)
                if item.recommendation:
                    rec = doc.add_paragraph(style="List Bullet 2")
                    rec.add_run(f"Рекомендация: {item.recommendation}").font.color.rgb = RGBColor(0x2A, 0x64, 0x96)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.post("/verify")
async def verify(req: VerifyRequest):
    report = await verify_grant_text(req.text, req.provider, req.model)
    return asdict(report)


@app.post("/export/docx")
async def export_docx(req: ExportRequest):
    content = _build_docx(req)
    filename = f"grantchecker_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/models")
async def models():
    return AVAILABLE_MODELS


@app.get("/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}


@app.get("/", response_class=HTMLResponse)
async def index():
    html_path = Path(__file__).parent / "templates" / "index.html"
    return HTMLResponse(content=html_path.read_text(encoding="utf-8"))
